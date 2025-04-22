import streamlit as st
import docx
from docx import Document
from datetime import datetime
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import io
import os
import pypandoc

# Set page config as the FIRST Streamlit command
st.set_page_config(page_title="Invoice Generator", page_icon="📄", layout="wide")

# Custom CSS for muted colors, rounded layout, and depth
st.markdown("""
    <style>
    /* General page styling */
    .stApp {
        background-color: #f0f2f6;
        font-family: 'Arial', sans-serif;
    }

    /* Header styling */
    h1 {
        color: #4a5e6a !important;
        text-align: center;
        text-shadow: 1px 1px 2px rgba(0, 0, 0, 0.1);
        padding: 20px;
        background: linear-gradient(145deg, #e6e9ef, #d5d9e0);
        border-radius: 15px;
        box-shadow: 5px 5px 15px rgba(0, 0, 0, 0.1), -5px -5px 15px rgba(255, 255, 255, 0.8);
        margin-bottom: 30px;
    }

    /* Section headers */
    h2 {
        color: #5a7d7c !important;
        margin-top: 20px;
        margin-bottom: 10px;
        padding: 10px 20px;
        background-color: #e2ece9;
        border-radius: 10px;
        box-shadow: 3px 3px 10px rgba(0, 0, 0, 0.1), -3px -3px 10px rgba(255, 255, 255, 0.7);
    }

    /* Form container styling */
    .stForm {
        background: linear-gradient(145deg, #e6e9ef, #d5d9e0);
        padding: 20px;
        border-radius: 15px;
        box-shadow: 5px 5px 15px rgba(0, 0, 0, 0.1), -5px -5px 15px rgba(255, 255, 255, 0.8);
        margin-bottom: 20px;
    }

    /* Input fields */
    .stTextInput > div > input,
    .stTextArea > div > textarea,
    .stNumberInput > div > input,
    .stDateInput > div > input {
        background-color: #f7f9fc !important;
        border: 1px solid #b0c4c3 !important;
        border-radius: 10px !important;
        padding: 10px !important;
        box-shadow: inset 2px 2px 5px rgba(0, 0, 0, 0.05), inset -2px -2px 5px rgba(255, 255, 255, 0.5) !important;
        color: #4a5e6a !important;
    }

    /* Buttons */
    .stButton > button {
        background: linear-gradient(145deg, #a3bffa, #7f9cfb);
        color: white !important;
        border: none !important;
        border-radius: 10px !important;
        padding: 10px 20px !important;
        box-shadow: 3px 3px 10px rgba(0, 0, 0, 0.1), -3px -3px 10px rgba(255, 255, 255, 0.7);
        transition: transform 0.1s ease-in-out;
    }

    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 5px 5px 15px rgba(0, 0, 0, 0.15), -5px -5px 15px rgba(255, 255, 255, 0.9);
    }

    /* Download buttons */
    .stDownloadButton > button {
        background: linear-gradient(145deg, #f4a261, #e76f51);
        color: white !important;
        border-radius: 10px !important;
        padding: 10px 20px !important;
        box-shadow: 3px 3px 10px rgba(0, 0, 0, 0.1), -3px -3px 10px rgba(255, 255, 255, 0.7);
    }

    /* Checkbox and labels */
    .stCheckbox > label {
        color: #5a7d7c !important;
    }

    /* Error and success messages */
    .stAlert {
        border-radius: 10px !important;
        box-shadow: 3px 3px 10px rgba(0, 0, 0, 0.1), -3px -3px 10px rgba(255, 255, 255, 0.7);
    }

    /* Items section */
    .stColumn {
        padding: 10px;
    }
    </style>
""", unsafe_allow_html=True)

# Reuse your existing classes and functions
class InvoiceData:
    def __init__(self):
        self.client_info = {}
        self.invoice_details = {}
        self.items = []
        self.financials = {}
        self.apply_late_fee = False
        self.invoice_number = ""

def format_currency(amount):
    if amount == 0:
        return ""
    elif amount == int(amount):
        return f"Rp {int(amount):,}"
    else:
        return f"Rp {amount:,.2f}"

def set_cell_border(cell, side, color="FFFFFF", sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    side_mapping = {
        'top': 'top', 'bottom': 'bottom', 'left': 'left', 'right': 'right'
    }
    border_name = side_mapping.get(side.lower())
    if border_name:
        border = parse_xml(f'<w:{border_name} {nsdecls("w")} w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>')
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
            tcPr.append(tcBorders)
        tcBorders.append(border)

def set_white_borders(cell, sz=4):
    for border in ['top', 'bottom', 'left', 'right']:
        set_cell_border(cell, border, color="FFFFFF", sz=sz)

def set_cell_font(cell, font_name="Courier New", font_size=10):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def apply_cell_style(cell, bg_color="#ddefd5"):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" />')
    cell._tc.get_or_add_tcPr().append(shading_elm)
    set_white_borders(cell, sz=6)
    set_cell_font(cell)

def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
    return doc

def update_items_table(doc, items):
    items_table = doc.tables[0]
    for i in range(len(items_table.rows)):
        for cell in items_table.rows[i].cells:
            set_white_borders(cell, sz=6)
    while len(items_table.rows) > 2:
        items_table._tbl.remove(items_table.rows[2]._tr)
    placeholder_row = items_table.rows[1]
    for item in items:
        row = items_table.add_row()
        row.cells[0].text = item['description']
        row.cells[1].text = format_currency(item['unit_price'])
        quantity = item['quantity']
        if quantity == int(quantity):
            row.cells[2].text = str(int(quantity))
        else:
            row.cells[2].text = str(quantity)
        row.cells[3].text = format_currency(item['total'])
        for i, cell in enumerate(row.cells):
            apply_cell_style(cell)
            alignments = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, 
                         WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignments[i]
    items_table._tbl.remove(placeholder_row._tr)
    return doc

def style_financial_table(doc, invoice_data):
    financial_table = doc.tables[1]
    for row in financial_table.rows:
        for cell in row.cells:
            set_white_borders(cell)
            set_cell_font(cell)
        for paragraph in row.cells[1].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if invoice_data.apply_late_fee:
        late_fee_cell = financial_table.rows[3].cells[0]
        if "LATE FEE" in late_fee_cell.text:
            original_text = late_fee_cell.text
            late_fee_cell.text = ""
            paragraph = late_fee_cell.paragraphs[0]
            run = paragraph.add_run(original_text)
            run.font.color.rgb = RGBColor.from_string('d95132')
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Courier New")

def get_next_invoice_number():
    count_file = "invoice_count.txt"
    year = "2025"
    if os.path.exists(count_file):
        with open(count_file, 'r') as f:
            try:
                count = int(f.read().strip())
            except ValueError:
                count = 0
    else:
        count = 0
    count += 1
    return f"INV{year}{count:03d}", count

def save_invoice_count(count):
    with open("invoice_count.txt", 'w') as f:
        f.write(str(count))

def validate_date_format(date_str):
    try:
        datetime.strptime(date_str, "%d.%m.%Y")
        return True
    except ValueError:
        return False

def generate_invoice(invoice_data):
    doc = Document('Invoice_Template_MarketixLab.docx')
    replacements = {**invoice_data.client_info, **invoice_data.invoice_details, **invoice_data.financials}
    if invoice_data.apply_late_fee:
        replacements['{{LATE FEE:}}'] = 'LATE FEE'
    else:
        replacements['{{LATE FEE:}}'] = ''
        replacements['[latefee]'] = ''
    doc = replace_placeholders(doc, replacements)
    doc = update_items_table(doc, invoice_data.items)
    style_financial_table(doc, invoice_data)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Courier New")
    
    # Generate DOCX
    docx_output = io.BytesIO()
    doc.save(docx_output)
    docx_output.seek(0)
    
    # Generate PDF using temporary files with pypandoc
    temp_docx = f"temp_{invoice_data.invoice_number}.docx"
    temp_pdf = f"temp_{invoice_data.invoice_number}.pdf"
    doc.save(temp_docx)
    
    # Convert DOCX to PDF using pypandoc
    pypandoc.convert_file(temp_docx, 'pdf', outputfile=temp_pdf)
    
    # Read PDF into BytesIO for download
    pdf_output = io.BytesIO()
    with open(temp_pdf, 'rb') as f:
        pdf_output.write(f.read())
    pdf_output.seek(0)
    
    # Clean up temporary files
    if os.path.exists(temp_docx):
        os.remove(temp_docx)
    if os.path.exists(temp_pdf):
        os.remove(temp_pdf)
    
    return (docx_output, f"Invoice_{invoice_data.invoice_number}.docx",
            pdf_output, f"Invoice_{invoice_data.invoice_number}.pdf")

# Streamlit App
st.title("📄 Invoice Generator")
st.markdown("Create professional invoices with ease using this streamlined tool.")

# Initialize session state for items and invoice date
if 'item_list' not in st.session_state:
    st.session_state.item_list = [{"description": "", "unit_price": 0.0, "quantity": 0.0}]

if 'use_today' not in st.session_state:
    st.session_state.use_today = True

if 'manual_invoice_date' not in st.session_state:
    st.session_state.manual_invoice_date = datetime.now()

# Client Information
st.header("Client Information")
with st.form(key="client_form"):
    client_name = st.text_input("Client Name", placeholder="Enter client name")
    client_phone = st.text_input("Client Phone", placeholder="Enter phone number")
    client_email = st.text_input("Client Email", placeholder="Enter email")
    client_address = st.text_area("Client Address", placeholder="Enter address")
    client_submit = st.form_submit_button("Save Client Info")

# Invoice Details
st.header("Invoice Details")
with st.form(key="invoice_form"):
    default_invoice_number, invoice_count = get_next_invoice_number()
    invoice_number = st.text_input("Invoice Number", value=default_invoice_number, help="Invoice number must start with 'INV2025'")
    
    # Invoice Date with manual input option using date picker
    st.session_state.use_today = st.checkbox("Use Today's Date", value=st.session_state.use_today, key="use_today_checkbox")
    if st.session_state.use_today:
        invoice_date = datetime.now().strftime("%d.%m.%Y")
        st.write(f"Invoice Date: {invoice_date}")
    else:
        st.session_state.manual_invoice_date = st.date_input(
            "Select Invoice Date",
            value=st.session_state.manual_invoice_date,
            key="manual_invoice_date_picker"
        )
        invoice_date = st.session_state.manual_invoice_date.strftime("%d.%m.%Y")
        st.write(f"Selected Invoice Date: {invoice_date}")
    
    # Due Date with date picker
    due_date_obj = st.date_input("Select Due Date", value=datetime.now(), key="due_date_picker")
    due_date = due_date_obj.strftime("%d.%m.%Y")
    
    invoice_submit = st.form_submit_button("Save Invoice Details")

# Items
st.header("Items")
# Ensure item_list is always a list
if not isinstance(st.session_state.item_list, list):
    st.warning("Item list was corrupted. Resetting to default.")
    st.session_state.item_list = [{"description": "", "unit_price": 0.0, "quantity": 0.0}]

# Functions to manage items
def add_item():
    st.session_state.item_list.append({"description": "", "unit_price": 0.0, "quantity": 0.0})

def remove_item(index):
    if len(st.session_state.item_list) > 1:
        st.session_state.item_list.pop(index)

# Display and edit items
for i in range(len(st.session_state.item_list)):
    col1, col2, col3, col4 = st.columns([3, 2, 2, 1])
    with col1:
        st.session_state.item_list[i]["description"] = st.text_input(
            f"Description {i+1}",
            value=st.session_state.item_list[i]["description"],
            key=f"desc_{i}"
        )
    with col2:
        st.session_state.item_list[i]["unit_price"] = st.number_input(
            f"Unit Price {i+1}",
            min_value=0.0,
            value=st.session_state.item_list[i]["unit_price"],
            key=f"price_{i}"
        )
    with col3:
        st.session_state.item_list[i]["quantity"] = st.number_input(
            f"Quantity {i+1}",
            min_value=0.0,
            value=st.session_state.item_list[i]["quantity"],
            key=f"qty_{i}"
        )
    with col4:
        if st.button("✕", key=f"delete_{i}"):
            remove_item(i)

# Add item button
st.button("Add Item", on_click=add_item)

# Financial Details
st.header("Financial Details")
with st.form(key="financial_form"):
    tax_rate = st.number_input("Tax Rate (%)", min_value=0.0, value=0.0, help="Enter tax rate as a percentage")
    discount = st.number_input("Discount Amount", min_value=0.0, value=0.0, help="Enter discount amount in Rp")
    apply_late_fee = st.checkbox("Apply Late Fee (2%)", value=False, help="Check to apply a 2% late fee")
    financial_submit = st.form_submit_button("Save Financial Details")

# Generate Invoice
if st.button("Generate Invoice"):
    try:
        # Validate inputs
        if not all([client_name, client_phone, client_email, client_address]):
            st.error("All client info fields are required")
        elif not all([invoice_number, invoice_date, due_date]):
            st.error("All invoice details are required")
        elif not invoice_number.startswith("INV2025"):
            st.error("Invoice number must start with 'INV2025'")
        elif not validate_date_format(invoice_date):
            st.error("Invoice date must be in the format dd.mm.yyyy (e.g., 21.04.2025)")
        elif not validate_date_format(due_date):
            st.error("Due date must be in the format dd.mm.yyyy (e.g., 28.04.2025)")
        elif not st.session_state.item_list or not any(item["description"] and item["unit_price"] > 0 and item["quantity"] > 0 for item in st.session_state.item_list):
            st.error("At least one valid item is required")
        else:
            invoice_data = InvoiceData()
            invoice_data.client_info = {
                '{{client_name}}': client_name,
                '{{client_phone}}': client_phone,
                '{{client_email}}': client_email,
                '{{client_address}}': client_address
            }
            invoice_data.invoice_details = {
                '{{invoice_number}}': invoice_number,
                '{{invoice_date}}': invoice_date,
                '{{due_date}}': due_date
            }
            invoice_data.items = [
                {
                    'description': item['description'],
                    'unit_price': item['unit_price'],
                    'quantity': item['quantity'],
                    'total': item['unit_price'] * item['quantity']
                } for item in st.session_state.item_list if item['description'] and item['unit_price'] > 0 and item['quantity'] > 0
            ]
            subtotal = sum(item['total'] for item in invoice_data.items)
            tax = subtotal * (tax_rate / 100)
            invoice_data.apply_late_fee = apply_late_fee
            late_fee = subtotal * 0.02 if apply_late_fee else 0
            total = subtotal + tax - discount + late_fee
            invoice_data.financials = {
                '[subtotal]': format_currency(subtotal),
                '[tax]': format_currency(tax),
                '[discount]': format_currency(discount),
                '[latefee]': format_currency(late_fee),
                '[grandtotal]': format_currency(total)
            }
            invoice_data.invoice_number = invoice_number
            docx_output, docx_filename, pdf_output, pdf_filename = generate_invoice(invoice_data)
            # Only update count if using auto-generated number
            if invoice_number == default_invoice_number:
                save_invoice_count(invoice_count)
            st.success(f"Invoice generated successfully!")
            col1, col2 = st.columns(2)
            with col1:
                st.download_button(
                    label="Download Invoice (DOCX)",
                    data=docx_output,
                    file_name=docx_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            with col2:
                st.download_button(
                    label="Download Invoice (PDF)",
                    data=pdf_output,
                    file_name=pdf_filename,
                    mime="application/pdf"
                )
    except Exception as e:
        st.error(f"An error occurred: {str(e)}")
